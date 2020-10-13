# -*- coding: utf-8 -*-
import time
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import pandas as pd
from spt_mysqlpy import plots
#import docx
#from cash_flow_report import get_date
from rep import rep

print('开始生成报告')

def insert_picture(cell, date, picture_name, picture_width, picture_height=None):
    temp_var = 0
    if date != "":
        run = cell.paragraphs[0].add_run(date)
        run.font.size = Pt(10)
        run.font.name = u'华文楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
        cell.add_paragraph()
        temp_var = 1
    para1 = cell.paragraphs[temp_var]
    para1.add_run().add_picture(doc_file_path + picture_name, width=picture_width, height=picture_height)
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER

#
doc_file_path = 'C:\\Users\\lenovo\\Desktop\\report\\'
document = Document(doc_file_path + 'report_demo.docx')

plots()
rep = rep()
document.paragraphs[0].clear()
run = document.paragraphs[0].add_run(("      "+rep[0]))
run.font.size = Pt(18)
run.font.name = u'华文楷体'
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
document.paragraphs[2].clear()
run = document.paragraphs[2].add_run(rep[1])
run.font.size = Pt(12)
run.font.name = u'华文楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
document.paragraphs[5].clear()
run = document.paragraphs[5].add_run(rep[2])
run.font.size = Pt(12)
run.font.name = u'华文楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')

#date = get_date()
#Date = [date[0], date[0], date[0], date[0],
#        date[1] + '至' + date[0], date[1] + '至' + date[0],
#        date[1] + '至' + date[0], date[1] + '至' + date[0],
#        date[2] + '至' + date[0], date[2] + '至' + date[0],
#        date[2] + '至' + date[0], date[2] + '至' + date[0]]
#picture_name = ['图1：前一交易日增减仓前五名品种持仓金额变化情况.png',
#                '图2：前一交易日成交金额前十名品种成交情况.png',
#                '图3：前一交易日各板块持仓金额变化情况.png',
#                '图4：前一交易日各板块成交金额情况况.png',
#                '图5：5日滚动增减仓前五名品种持仓金额变化情况.png',
#                '图6：5日滚动成交金额前十名品种成交情况.png',
#                '图7：5日滚动各板块持仓金额变化情况.png',
#                '图8：5日滚动各板块成交金额情况.png',
#                '图9：10日滚动增减仓前五名品种持仓金额变化情况.png',
#                '图10：10日滚动成交金额前十名品种成交情况.png',
#                '图11：10日滚动各板块持仓金额变化情况.png',
#                '图12：10日滚动各板块成交金额情况.png',
#                '图13： 各版块持仓金额(亿元).png',
#                '图14： 各版块成交金额(亿元).png',
#                '图15：中国商品期货市场持仓金额、成交金额与Wind商品指数.png',
#                '图16：有色板块持仓金额、成交金额与Wind有色指数.png',
#                '图17：油脂油料板块持仓金额、成交金额与Wind油脂油料指数.png',
#                '图18：软商品板块持仓金额、成交金额与Wind软商品指数.png',
#                '图19：农副产品板块持仓金额、成交金额与Wind农副产品指数.png',
#                '图20：谷物板块持仓金额、成交金额与Wind谷物指数.png',
#                '图21：贵金属板块持仓金额、成交金额与Wind贵金属指数.png',
#                '图22：能源板块持仓金额、成交金额与Wind能源指数.png',
#                '图23：煤焦钢矿持仓金额、成交金额与Wind煤焦钢矿指数.png',
#                '图24：非金属建材持仓金额、成交金额与Wind非金属建材指数.png',
#                '图25：化工板块持仓金额、成交金额与Wind化工指数.png',
#                '图26：有色板块月度成交金额(亿元).png',
#                '图27：油脂油料板块月度成交金额(亿元).png',
#                '图28：软商品板块月度成交金额(亿元).png',
#                '图29：农副产品板块月度成交金额(亿元).png',
#                '图30：谷物板块月度成交金额(亿元).png',
#                '图31：贵金属板块月度成交金额(亿元).png',
#                '图32：能源板块月度成交金额(亿元).png',
#                '图33：煤焦钢矿板块月度成交金额(亿元).png',
#                '图34：非金属建材板块月度成交金额(亿元).png',
#                '图35：化工板块月度成交金额(亿元).png']
picture_name = ['spt1.png',
               'spt2.png',
                'spt3.png',
                'spt4.png',
                'spt2_1.png',
               'spt2_2.png',
                'spt2_3.png',
                'spt2_4.png',
                'spt2_5.png',
               'spt2_6.png',
                'IF_long.png',
                'IF_short.png',
                'IC_long.png',
                'IC_short.png',
                'IH_long.png',
                'IH_short.png',
                'TS_long.png',
                'TS_short.png',
                'TF_long.png',
                'TF_short.png',
                'T_long.png',
                'T_short.png',
                'IF_basis.png',
                'IC_basis.png',
                'IH_basis.png']
#            
#                '图26：有色板块月度成交金额(亿元).png',
#                '图27：油脂油料板块月度成交金额(亿元).png',
#                '图28：软商品板块月度成交金额(亿元).png',
#                '图29：农副产品板块月度成交金额(亿元).png',
#                '图30：谷物板块月度成交金额(亿元).png',
#                '图31：贵金属板块月度成交金额(亿元).png',
#                '图32：能源板块月度成交金额(亿元).png',
#                '图33：煤焦钢矿板块月度成交金额(亿元).png',
#                '图34：非金属建材板块月度成交金额(亿元).png',
#                '图35：化工板块月度成交金额(亿元).png']
#for i in range(0, 6):
#    table = document.tables[i]
#    insert_picture(table.cell(1, 0), Date[i * 2], picture_name[i * 2], Inches(3.4))
#    insert_picture(table.cell(1, 2), Date[i * 2 + 1], picture_name[i * 2 + 1], Inches(3.4))
#
for m in range(0, 5):
    table = document.tables[m]
    insert_picture(table.cell(1, 0), "", picture_name[m*2], Inches(3.4))
    insert_picture(table.cell(1, 2), "", picture_name[m*2 + 1], Inches(3.4))
    
    
df=pd.read_csv('rank.csv')[['合计','多单量','多单占比(%)','空单量','空单占比(%)','净多单','净空单']]
#print(df)
#
table = document.tables[5]
for row in range(18):
    for col in range(6):
        table.cell(row+1,col+2).text=str(df.loc[row,df.columns.values[col+1]])
        run=table.cell(row+1,col+2).paragraphs[0].runs
        run[0].font.size = Pt(10)
        run[0].font.name = 'Times New Roman'
        run[0].font.color.rgb = RGBColor(0,0,0)
        
#        

    

for m in range(6, 13):
    table = document.tables[m]
    insert_picture(table.cell(1, 0), "", picture_name[m*2-2], Inches(3.4))
    insert_picture(table.cell(1, 2), "", picture_name[m*2 - 1], Inches(3.4))
    
table = document.tables[13]
insert_picture(table.cell(1, 0), "", picture_name[24], Inches(3.4))


def next_weekday(date, *args): # 获取下一个工作日日期的字符串
    if args: # 如果因为节假日导致下一个工作日发生变化，手动输入下一个工作日的日期
        return str(args[0])
    else:
        dayofweek = date.isoweekday()
        if dayofweek < 5:
            num_days = 1
        else:
            num_days = 8 - dayofweek
        return date + datetime.timedelta(days = num_days)

def revise_header_footer(file, next_day): # 修改页眉页脚
    # 修改页眉日期
    header_para = file.sections[0].first_page_header.paragraphs[2]
    [_, date1] = header_para.text.rsplit(" ", 1)
    header_runs = header_para.runs
    for i in header_runs:
        if i.text.replace(" ", "") != '' and i.text.replace(" ", "") in date1:
            i.text = ''
    header_runs[-1].text = next_day
    # 修改页脚日期
    footer_para = file.sections[0].footer.paragraphs[0]
    [date2, _] = footer_para.text.split(" ", 1)
    footer_runs = footer_para.runs
    for i in footer_runs:
        if i.text.replace(" ", "") != '' and i.text.replace(" ", "") in date2:
            i.text = ''
    footer_runs[0].text = next_day
    return file

def get_file_name(next_day, title):
    the_month = str(next_day.month)
    the_day = str(next_day.day)
    if len(the_month) < 2:
        the_month = '0' + the_month
    if len(the_day) < 2:
        the_day = '0' + the_day
    the_date = str(next_day.year) + the_month + the_day
    return "金融期货市场流动性日报" + the_date + '：' + title + '.docx'

next_day = next_weekday(datetime.date.today())
document = revise_header_footer(document, str(next_day))
file_name = get_file_name(next_day, rep[0])
document.save(doc_file_path + file_name)




#    
#
document.save(doc_file_path +"商品期货市场流动性日报"+ time.strftime('%Y%m%d', time.localtime(time.time())) +'：' + rep[0] + '.docx')
print('报告生成成功，地址：' + doc_file_path)